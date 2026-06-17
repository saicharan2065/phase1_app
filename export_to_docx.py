import os
import docx

def export_project_to_docx():
    doc = docx.Document()
    doc.add_heading('Antigravity OS - Source Code', 0)
    
    target_dirs = ['tabs', 'agents', 'data']
    
    for root, dirs, files in os.walk('.'):
        # Skip venv, cache, git, etc.
        if any(ignore in root for ignore in ['.git', '__pycache__', 'venv', '.env', 'storage']):
            continue
            
        for file in files:
            if file.endswith('.py') or file.endswith('.md') or file == 'requirements.txt':
                file_path = os.path.join(root, file)
                
                # Add File Heading
                doc.add_heading(file_path, level=1)
                
                # Read content
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Add code block
                    p = doc.add_paragraph(content)
                    p.style.font.name = 'Courier New'
                except Exception as e:
                    doc.add_paragraph(f"Error reading file: {str(e)}")
                    
    doc.save('Project_Source_Code.docx')
    print("Successfully exported all code to Project_Source_Code.docx")

if __name__ == "__main__":
    export_project_to_docx()
